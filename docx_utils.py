# docx_utils.py
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def read_docx_content(file_io):
    try:
        doc = docx.Document(file_io)
        return '\n'.join([para.text for para in doc.paragraphs])
    except Exception as e:
        print(f"Error reading DOCX content: {e}")
        raise

def create_feedback_doc(student_name, feedback, original_essay):
    try:
        if not feedback:
            raise ValueError("Feedback content is empty or None")

        # Create a new Document
        doc = docx.Document()

        # Adjust standard spacing for the "Normal" paragraph style
        style = doc.styles['Normal']
        if style.type == docx.enum.style.WD_STYLE_TYPE.PARAGRAPH: # type: ignore
            paragraph_format = style.paragraph_format # type: ignore
            paragraph_format.space_after = Pt(6)

            # Set the font if applicable
            if style.font is not None: # type: ignore
                style.font.name = 'Times New Roman' # type: ignore
                style.font.size = Pt(12) # type: ignore

        # Add student's name as the title (centered)
        title = doc.add_heading(student_name, level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        def parse_markdown(text, paragraph):
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = paragraph.add_run(part[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(part)

        # Add the original essay with proper line breaks and spacing
        doc.add_heading("Ursprunglig uppsats", level=1)
        paragraphs = re.split(r'\n\n|\n', original_essay)  # Split on double or single newlines
        for para in paragraphs:
            if para.strip():  # Only add non-empty paragraphs
                p = doc.add_paragraph()
                parse_markdown(para.strip(), p)
                if p.style and hasattr(p.style, 'paragraph_format') and p.style.paragraph_format:
                    p.style.paragraph_format.space_after = Pt(6)

        # Add space after the original essay
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

        # Split the feedback into sections
        sections = re.split(r'##\s*', feedback)

        for section in sections:
            if not section.strip():
                continue
            lines = section.strip().split('\n')
            heading = lines[0].strip()
            content = '\n'.join(lines[1:]).strip()

            # Add section heading for all main sections
            if heading.lower() in ["feedback", "förtjänster", "förbättringsområden", "redaktörens förslag före publicering"]:
                doc.add_heading(heading, level=1)

            if heading.lower() == "redaktörens förslag före publicering":
                # Add the edited essay content with proper spacing and title handling
                paragraphs = re.split(r'\n\n|\n', content)  # Split on double or single newlines
                for para in paragraphs:
                    stripped_para = para.strip()
                    if stripped_para:
                        # Check for markdown title levels
                        title_match = re.match(r'^(#{1,6})\s+(.+)$', stripped_para)
                        if title_match:
                            level = len(title_match.group(1))  # Number of # determines the level
                            title_text = title_match.group(2)
                            doc.add_heading(title_text, level=min(level, 9))  # Word supports up to 9 levels
                        else:
                            p = doc.add_paragraph()
                            parse_markdown(stripped_para, p)
                            if p.style and hasattr(p.style, 'paragraph_format') and p.style.paragraph_format:
                                p.style.paragraph_format.space_after = Pt(6)
                continue

            # Process content for other sections
            paragraphs = re.split(r'\n\n|\n', content)  # Split on double or single newlines
            for i, para in enumerate(paragraphs):
                stripped_para = para.strip()
                if stripped_para:
                    if i == 0 and heading.lower() == "feedback":
                        # For the first paragraph of the feedback section (greeting)
                        p = doc.add_paragraph()
                        parse_markdown(stripped_para, p)
                        if p.style and hasattr(p.style, 'paragraph_format') and p.style.paragraph_format:
                            p.style.paragraph_format.space_after = Pt(0)  # No extra space after greeting
                    elif re.match(r'^\d+\.|\*', stripped_para):
                        # For numbered or bullet points
                        lines = stripped_para.split('\n')
                        for line in lines:
                            text = re.sub(r'^\d+\.|\*\s*', '', line.strip())
                            p = doc.add_paragraph(style='List Bullet' if line.strip().startswith('*') else 'List Number')
                            parse_markdown(text, p)
                    else:
                        # For regular paragraphs
                        p = doc.add_paragraph()
                        parse_markdown(stripped_para, p)
                        if p.style and hasattr(p.style, 'paragraph_format') and p.style.paragraph_format:
                            p.style.paragraph_format.space_after = Pt(6)

        print("Feedback-DOCX document created.")
        return doc
    except Exception as e:
        print(f"Error creating feedback DOCX: {e}")
        raise